"""
Document and YouTube ingestion utilities.

This is a direct port of PART 3 ("INGEST YOUR DATA") from the original
RAD_RAG_Sandbox notebook, with all Colab-only calls (`google.colab.files`,
`!apt-get`, `!pip install`) removed. System packages (tesseract-ocr,
poppler-utils) are installed at Docker build time instead — see the
Dockerfile.
"""
import hashlib
import logging
import re
import shutil
from pathlib import Path
from urllib.parse import urlparse

import pdfplumber
import pytesseract
import trafilatura
from docx import Document as DocxDocument
from llama_index.core import Document
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api.formatters import TextFormatter

from app.config import settings
from app.core.state_store import register_doc_source

logger = logging.getLogger("docmind.ingestion")
ytt_api = YouTubeTranscriptApi()

class IngestionError(Exception):
    """Raised for any recoverable ingestion failure (bad file, bad URL, etc.)."""


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def extract_video_id(url: str) -> str:
    """Extract the video ID from common YouTube URL formats."""
    patterns = [
        r"v=([^&]+)",
        r"youtu\.be/([^?]+)",
        r"embed/([^?]+)",
        r"shorts/([^?]+)",
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    raise IngestionError(f"Invalid YouTube URL: {url}")


def clean_text(text: str) -> str:
    """Removes binary artifacts and drops text that looks like garbage/OCR noise."""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    alnum_ratio = len(re.findall(r"[a-zA-Z0-9\s\.\,\;\:\-\"\'\?\!]", text)) / max(1, len(text))
    return text if alnum_ratio > 0.6 else ""


def extract_pdf_text(file_path: str) -> str:
    """Attempts text extraction via pdfplumber, falls back to OCR (tesseract) if needed."""
    full_text = ""

    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    cleaned = clean_text(page_text)
                    if cleaned:
                        full_text += cleaned + "\n"
    except Exception as e:
        logger.warning("pdfplumber error on %s: %s", file_path, e)

    if len(full_text.strip()) < 50:
        logger.info("Low text quality from %s (%d chars). Attempting OCR...", file_path, len(full_text))
        try:
            from pdf2image import convert_from_path

            images = convert_from_path(file_path, dpi=300)
            for img in images:
                ocr_text = pytesseract.image_to_string(img)
                cleaned = clean_text(ocr_text)
                if cleaned:
                    full_text += cleaned + "\n"
            logger.info("OCR extracted %d characters from %s.", len(full_text), file_path)
        except Exception as e:
            logger.error("OCR failed for %s: %s", file_path, e)

    if clean_text(full_text) == "":
        return ""

    return full_text


def detect_resource(resource: str) -> str:
    """Detect whether the input is a YouTube URL, a generic URL, a local document, or unknown."""
    resource = resource.strip()

    if resource.startswith(("http://", "https://")):
        domain = urlparse(resource).netloc.lower()
        if "youtube.com" in domain or "youtu.be" in domain:
            return "youtube"
        return "url"

    path = Path(resource)
    if path.exists() and path.suffix.lower() in settings.supported_extensions:
        return "document"

    return "unknown"


# ---------------------------------------------------------------------------
# Single-document loaders (used both for ingestion endpoints and full reindex)
# ---------------------------------------------------------------------------

def load_single_document(file_path: Path) -> Document | None:
    """Parses one file on disk into a llama-index Document, or None if unreadable."""
    ext = file_path.suffix.lower()

    if ext == ".pdf":
        text = extract_pdf_text(str(file_path))
        if not text:
            logger.warning("Skipping %s - completely unreadable.", file_path.name)
            return None
        return Document(text=text, metadata={"source": file_path.name, "type": "pdf"})

    if ext in (".txt", ".md"):
        text = file_path.read_text(encoding="utf-8", errors="ignore")
        return Document(text=text, metadata={"source": file_path.name, "type": "text"})

    if ext == ".docx":
        docx = DocxDocument(str(file_path))
        text = "\n".join(p.text for p in docx.paragraphs if p.text.strip())
        return Document(text=text, metadata={"source": file_path.name, "type": "docx"})

    logger.warning("Unsupported format: %s", ext)
    return None


def load_documents_robust(folder_path: str | None = None) -> list[Document]:
    """Walks the user_docs folder and parses every supported file. Used for full reindex."""
    folder = Path(folder_path or settings.user_docs_dir)
    documents: list[Document] = []

    for file in sorted(folder.iterdir()):
        if not file.is_file():
            continue
        if file.suffix.lower() not in settings.supported_extensions:
            continue

        logger.info("Processing: %s", file.name)
        try:
            doc = load_single_document(file)
            if doc is not None:
                documents.append(doc)
        except Exception as e:
            logger.error("Error processing %s: %s", file.name, e)

    return documents


# ---------------------------------------------------------------------------
# Ingestion entry points (called from the API routers)
# ---------------------------------------------------------------------------

def save_uploaded_file(filename: str, content: bytes) -> Path:
    """Saves raw uploaded bytes into the user_docs directory."""
    ext = Path(filename).suffix.lower()
    if ext not in settings.supported_extensions:
        raise IngestionError(
            f"Unsupported file type '{ext}'. Supported types: {', '.join(settings.supported_extensions)}"
        )

    destination = Path(settings.user_docs_dir) / filename
    if destination.exists():
        raise IngestionError(f"A document named '{filename}' already exists.")

    destination.write_bytes(content)
    register_doc_source(filename, source_type="upload")
    return destination


def ingest_local_path(path: str) -> Path:
    """Copies an existing file on disk into the user_docs directory (non-upload use case)."""
    source = Path(path)
    if not source.exists():
        raise IngestionError(f"File not found: {path}")

    destination = Path(settings.user_docs_dir) / source.name
    if destination.exists():
        raise IngestionError(f"A document named '{source.name}' already exists.")

    shutil.copy2(source, destination)
    register_doc_source(source.name, source_type="upload")
    return destination


def fetch_youtube_transcript(url: str) -> Path:
    """Downloads a YouTube transcript and saves it as a .txt file in user_docs."""
    video_id = str(extract_video_id(url))
    # print("youtube video id: ", video_id)
    transcript_path = Path(settings.user_docs_dir) / f"{video_id}.txt"

    if transcript_path.exists():
        raise IngestionError(f"Transcript for video '{video_id}' already exists.")

    logger.info("Fetching transcript for video_id=%s", video_id)
    try:
        # transcript_text = "this testing text" # testing purpose
        transcript = ytt_api.fetch(video_id=video_id)
        transcript_text = TextFormatter().format_transcript(transcript)
    except Exception as e:
        raise IngestionError(f"Could not fetch transcript: {e}") from e

    transcript_path.write_text(transcript_text, encoding="utf-8")
    register_doc_source(transcript_path.name, source_type="youtube", source_url=url)
    return transcript_path


def fetch_web_page(url: str) -> Path:
    """Downloads a web page and extracts its main article text (boilerplate/nav/ads stripped)."""
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https") or not parsed.netloc:
        raise IngestionError(f"Invalid URL: {url}")

    domain = re.sub(r"[^a-zA-Z0-9]+", "-", parsed.netloc).strip("-")
    url_hash = hashlib.sha1(url.encode("utf-8")).hexdigest()[:10]
    filename = f"web_{domain}_{url_hash}.txt"
    destination = Path(settings.user_docs_dir) / filename

    if destination.exists():
        raise IngestionError(f"This page has already been ingested ('{filename}').")

    logger.info("Fetching web page: %s", url)
    try:
        downloaded = trafilatura.fetch_url(url)
        if not downloaded:
            raise IngestionError("Could not download the page (network error or blocked).")

        extracted = trafilatura.extract(
            downloaded, include_comments=False, include_tables=True, favor_recall=True
        )
    except IngestionError:
        raise
    except Exception as e:
        raise IngestionError(f"Failed to fetch/parse page: {e}") from e

    cleaned = clean_text(extracted or "")
    if not cleaned or len(cleaned.strip()) < 50:
        raise IngestionError("The page didn't contain enough readable text to ingest.")

    destination.write_text(f"Source: {url}\n\n{cleaned}", encoding="utf-8")
    register_doc_source(filename, source_type="web", source_url=url)
    return destination
